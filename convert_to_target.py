# -*- coding: utf-8 -*-
"""
Markdown转Word文档转换脚本
用于将AI IDE与CLI工具定价调研报告转换为Word格式
"""

import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_chinese_font(run, font_name='微软雅黑', font_size=11):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._element
    rFonts = r.find(qn('w:rPr'))
    if rFonts is None:
        rFonts = OxmlElement('w:rPr')
        r.insert(0, rFonts)
    rFonts_elem = rFonts.find(qn('w:rFonts'))
    if rFonts_elem is None:
        rFonts_elem = OxmlElement('w:rFonts')
        rFonts.insert(0, rFonts_elem)
    rFonts_elem.set(qn('w:eastAsia'), font_name)


def add_heading(doc, text, level):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_chinese_font(run, '微软雅黑', 16 - level if level <= 3 else 11)
    return heading


def add_paragraph(doc, text, bold=False):
    """添加段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_chinese_font(run, '微软雅黑', 11)
    if bold:
        run.bold = True
    return para


def parse_table(lines, start_idx):
    """解析Markdown表格"""
    table_data = []
    i = start_idx
    while i < len(lines) and '|' in lines[i]:
        row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
        if row and not all(c.replace('-', '').replace(':', '') == '' for c in row):
            table_data.append(row)
        i += 1
    return table_data, i


def add_table(doc, table_data):
    """添加表格"""
    if not table_data:
        return
    
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, '微软雅黑', 10)
                    if i == 0:
                        run.bold = True


def convert_md_to_docx(md_file, docx_file):
    """将Markdown文件转换为Word文档"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    doc = Document()
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        if line.startswith('# '):
            add_heading(doc, line[2:], 1)
        elif line.startswith('## '):
            add_heading(doc, line[3:], 2)
        elif line.startswith('### '):
            add_heading(doc, line[4:], 3)
        elif line.startswith('#### '):
            add_heading(doc, line[5:], 4)
        elif line.startswith('|'):
            table_data, i = parse_table(lines, i)
            add_table(doc, table_data)
            continue
        elif line.startswith('**') and line.endswith('**'):
            add_paragraph(doc, line[2:-2], bold=True)
        elif line.startswith('- ') or line.startswith('* '):
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(line[2:])
            set_chinese_font(run, '微软雅黑', 11)
        elif line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            para = doc.add_paragraph()
            run = para.add_run('\n'.join(code_lines))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif line.startswith('---'):
            doc.add_paragraph()
        else:
            add_paragraph(doc, line)
        
        i += 1
    
    doc.save(docx_file)
    print(f"Word文档已保存到: {docx_file}")


if __name__ == '__main__':
    md_file = r"a:\project\the_ai_math_teacher_revised - cn\AI_IDE与CLI工具定价及廉价使用顶级大模型策略调研报告_2026-03-18.md"
    docx_file = r"C:\Users\35928\Documents\研究资料\市场调研报告\AI_IDE与CLI工具定价及廉价使用顶级大模型策略调研报告_2026-03-18.docx"
    convert_md_to_docx(md_file, docx_file)
